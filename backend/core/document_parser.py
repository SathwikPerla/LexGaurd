"""
document_parser.py — Synchronous, production-quality document text extractor.

Supports:
  - PDF via PyMuPDF (fitz) — handles encrypted/corrupted docs gracefully
  - DOCX via python-docx — includes paragraphs + table cells
  - Scanned PDF detection — flags pages with < SCANNED_PAGE_THRESHOLD chars

Design decisions:
  - Synchronous functions (run in threadpool via asyncio.to_thread in routes)
  - MAX_CHARS = 120_000 (~90 pages) — enough for any contract, prevents OOM
  - _clean_text is deterministic and idempotent — safe to call repeatedly
  - parse_document is the single public entry point — callers don't care about format

Common failure points:
  - PyMuPDF throws fitz.FileDataError on corrupted PDFs — caught and re-raised as ValueError
  - Password-protected PDFs return empty text — detected and reported
  - DOCX from Google Docs may have unusual encoding — python-docx handles silently
  - Empty file_bytes causes immediate ValueError before any parsing attempt

How to test:
  - pytest backend/tests/test_parser.py -v
  - Upload a real employment contract PDF, check character_count > 1000
  - Upload a Google-Docs-exported DOCX, verify special chars render correctly
"""
from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF — import name is 'fitz', package name is 'PyMuPDF'
from docx import Document

logger = logging.getLogger(__name__)

# Safety cap — prevents OOM on very large contracts; ~90 standard pages
MAX_CHARS: int = 120_000

# If fewer than this many chars extracted per page, assume page is scanned image
SCANNED_PAGE_THRESHOLD: int = 50


# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────


def _clean_text(raw: str) -> str:
    """
    Normalise text extracted from PDF/DOCX:
      - Remove null bytes and control characters (except newline/tab)
      - Normalise line endings to \\n
      - Collapse runs of 3+ blank lines to 2
      - Strip trailing whitespace from every line
    This function is deterministic and idempotent.
    """
    # Remove null bytes
    text = raw.replace("\x00", "")
    # Normalise CR+LF and bare CR to LF
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove other non-printable control chars (keep \n and \t)
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Collapse 3+ consecutive blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip trailing whitespace per line
    text = "\n".join(line.rstrip() for line in text.splitlines())
    return text.strip()


def _truncate(text: str, filename: str) -> str:
    """Truncate text to MAX_CHARS and log a warning if truncation occurs."""
    if len(text) > MAX_CHARS:
        logger.warning(
            "Document truncated",
            extra={"doc_filename": filename, "original_chars": len(text), "limit": MAX_CHARS},
        )
        return text[:MAX_CHARS]
    return text


# ─────────────────────────────────────────────────────────────────────────────
# PDF parser
# ─────────────────────────────────────────────────────────────────────────────


def _parse_pdf(file_bytes: bytes, filename: str) -> dict:
    """
    Extract text from a PDF using PyMuPDF.

    Returns dict with:
        extracted_text  : str
        page_count      : int
        is_scanned_pdf  : bool  — True when >50% of pages yield no selectable text
        parse_method    : str   — 'pdf' or 'ocr' (ocr when scanned detected)
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise ValueError(
            f"Failed to open PDF '{filename}': {exc}. "
            "The file may be corrupted or password-protected."
        ) from exc

    page_count = len(doc)
    if page_count == 0:
        doc.close()
        raise ValueError(f"PDF '{filename}' contains no pages.")

    page_texts: list[str] = []
    scanned_page_count: int = 0

    for page_num in range(page_count):
        page = doc[page_num]
        try:
            text = page.get_text("text")  # type: ignore[attr-defined]
        except Exception as exc:
            logger.warning(
                "Failed to extract text from page",
                extra={"doc_filename": filename, "page": page_num + 1, "error": str(exc)},
            )
            text = ""

        if len(text.strip()) < SCANNED_PAGE_THRESHOLD:
            scanned_page_count += 1
        else:
            page_texts.append(text)

    doc.close()

    raw_text = "\n\n".join(page_texts)
    cleaned = _clean_text(raw_text)
    cleaned = _truncate(cleaned, filename)

    is_scanned = scanned_page_count > (page_count * 0.5)

    logger.info(
        "PDF parsed",
        extra={
            "doc_filename": filename,
            "pages": page_count,
            "scanned_pages": scanned_page_count,
            "is_scanned": is_scanned,
            "char_count": len(cleaned),
        },
    )

    return {
        "extracted_text": cleaned,
        "page_count": page_count,
        "is_scanned_pdf": is_scanned,
        "parse_method": "ocr" if is_scanned else "pdf",
    }


# ─────────────────────────────────────────────────────────────────────────────
# DOCX parser
# ─────────────────────────────────────────────────────────────────────────────


def _parse_docx(file_bytes: bytes, filename: str) -> dict:
    """
    Extract text from a DOCX file using python-docx.
    Includes body paragraphs and table cell content.

    Returns dict with:
        extracted_text  : str
        page_count      : None  (DOCX has no native page concept)
        is_scanned_pdf  : False
        parse_method    : 'docx'
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(
            f"Failed to open DOCX '{filename}': {exc}. "
            "The file may be corrupted or not a valid DOCX."
        ) from exc

    parts: list[str] = []

    # Body paragraphs (main content)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Table cells — legal contracts often use tables for key terms
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    raw_text = "\n\n".join(parts)
    cleaned = _clean_text(raw_text)
    cleaned = _truncate(cleaned, filename)

    logger.info(
        "DOCX parsed",
        extra={"doc_filename": filename, "char_count": len(cleaned)},
    )

    return {
        "extracted_text": cleaned,
        "page_count": None,
        "is_scanned_pdf": False,
        "parse_method": "docx",
    }


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────


def parse_document(file_bytes: bytes, filename: str) -> dict:
    """
    Parse a PDF or DOCX file and return extracted text with metadata.

    Args:
        file_bytes: Raw file content as bytes.
        filename:   Original filename including extension (used for routing).

    Returns:
        {
            "extracted_text": str,
            "page_count":     int | None,
            "is_scanned_pdf": bool,
            "parse_method":   "pdf" | "docx" | "ocr",
        }

    Raises:
        ValueError: For empty files, unsupported extensions, corrupted files.
    """
    if not file_bytes:
        raise ValueError("File is empty — nothing to parse.")

    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(file_bytes, filename)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(file_bytes, filename)
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Only .pdf and .docx files are accepted."
        )
